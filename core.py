from __future__ import annotations

import os
import re
import shutil
import subprocess
import sys
import unicodedata
import tempfile
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Optional

import pandas as pd
import pdfplumber
try:
    import pymupdf as fitz
except ImportError:
    try:
        import fitz  # type: ignore[no-redef]
    except ImportError:
        fitz = None

try:
    import pytesseract
except ImportError:
    pytesseract = None

from PIL import Image, ImageFilter, ImageOps
from build_config import OCR_ENABLED
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

TABLE_COLUMNS = [
    "Α/Α PDF",
    "Μητρώο",
    "Επώνυμο",
    "Όνομα",
    "Πατρώνυμο",
    "Βαθμός",
    "Οργανική",
    "Τρόπος Ταύτισης",
]


# ------------------------------
# Βοηθητικές συναρτήσεις κοινού σκοπού
# ------------------------------

def app_base_dir() -> Path:
    """Επιστρέφει τον φάκελο βάσης τόσο σε script όσο και σε PyInstaller build."""
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return Path(sys._MEIPASS)
    return Path(__file__).resolve().parent


def is_ocr_build() -> bool:
    """Δείχνει αν τρέχει η OCR έκδοση του προγράμματος."""
    return bool(OCR_ENABLED)


def require_ocr_capability() -> None:
    """
    Σταματά την OCR διαδρομή όταν τρέχει η Lite έκδοση
    ή όταν λείπει το PyMuPDF που χρειάζεται για rasterize του PDF.
    """
    if not OCR_ENABLED:
        raise RuntimeError(
            "Το PDF δεν έδωσε αναγνώσιμο text layer. "
            "Χρησιμοποίησε την OCR έκδοση του προγράμματος."
        )

    if fitz is None:
        raise RuntimeError(
            "Η OCR έκδοση δεν έχει το PyMuPDF (pymupdf). "
            "Χρειάζεται για μετατροπή του PDF σε εικόνα πριν το OCR."
        )


_TESSERACT_RUNTIME_READY = False
_TESSERACT_RUNTIME_DIR: Optional[Path] = None
_TESSERACT_RUNTIME_MODE = ""


def _native_path(path: Path) -> str:
    """Επιστρέφει απόλυτο path με native separators του λειτουργικού."""
    return str(path.resolve())


def _set_tessdata_prefix(value: Optional[Path]) -> None:
    if value is None:
        os.environ.pop("TESSDATA_PREFIX", None)
        return
    os.environ["TESSDATA_PREFIX"] = _native_path(value)


def _probe_tesseract_runtime(tesseract_cmd: str, required: list[str]) -> tuple[Optional[Path], str]:
    """
    Δοκιμάζει πραγματικά τον bundled Tesseract με --list-langs και κρατά
    μόνο το combo env/argument που όντως δουλεύει στο τρέχον runtime.
    """
    cmd_path = Path(tesseract_cmd).resolve()
    tesseract_dir = cmd_path.parent
    tessdata_dir = (tesseract_dir / "tessdata").resolve()

    if not tessdata_dir.is_dir():
        raise RuntimeError(f"Δεν βρέθηκε φάκελος tessdata δίπλα στο Tesseract: {tessdata_dir}")

    missing_files = [lang for lang in required if not (tessdata_dir / f"{lang}.traineddata").is_file()]
    if missing_files:
        raise RuntimeError(
            "Λείπουν language packs από το bundled tessdata: "
            f"{', '.join(missing_files)}. Τοποθεσία: {tessdata_dir}"
        )

    trials: list[tuple[str, Optional[Path], bool]] = [
        ("arg_only", None, True),
        ("env_tessdata_only", tessdata_dir, False),
        ("env_parent_only", tesseract_dir, False),
        ("env_tessdata_plus_arg", tessdata_dir, True),
        ("env_parent_plus_arg", tesseract_dir, True),
    ]

    last_output = ""
    for mode, env_prefix, use_arg in trials:
        env = os.environ.copy()
        if env_prefix is None:
            env.pop("TESSDATA_PREFIX", None)
        else:
            env["TESSDATA_PREFIX"] = _native_path(env_prefix)

        cmd = [tesseract_cmd]
        if use_arg:
            cmd.extend(["--tessdata-dir", _native_path(tessdata_dir)])
        cmd.append("--list-langs")

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            env=env,
            errors="replace",
        )
        combined = f"{result.stdout}\n{result.stderr}".strip()
        last_output = combined
        listed = {line.strip() for line in combined.splitlines() if line.strip()}
        if result.returncode == 0 and all(lang in listed for lang in required):
            return (tessdata_dir if use_arg else None, mode)

    raise RuntimeError(
        "Ο bundled Tesseract βρέθηκε αλλά δεν μπορεί να αρχικοποιήσει τις γλώσσες OCR. "
        f"Tesseract: {cmd_path} | Tessdata: {tessdata_dir} | Τελευταία έξοδος: {last_output}"
    )


def configure_tesseract() -> str:
    """
    Ρυθμίζει αυτόματα το path του Tesseract.

    Σειρά αναζήτησης:
    1. Bundled binary μέσα στο project / exe
    2. Κλασική εγκατάσταση Windows
    3. Ό,τι βρίσκεται ήδη στο PATH
    """
    require_ocr_capability()

    base_dir = app_base_dir()

    candidates = [
        base_dir / "third_party" / "tesseract" / "tesseract.exe",
        Path(__file__).resolve().parent / "third_party" / "tesseract" / "tesseract.exe",
        Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe"),
        Path(r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"),
        Path("/opt/homebrew/bin/tesseract"),
        Path("/usr/local/bin/tesseract"),
    ]

    which_path = shutil.which("tesseract")
    if which_path:
        candidates.append(Path(which_path))

    for candidate in candidates:
        if candidate.is_file():
            pytesseract.pytesseract.tesseract_cmd = str(candidate.resolve())

            global _TESSERACT_RUNTIME_READY, _TESSERACT_RUNTIME_DIR, _TESSERACT_RUNTIME_MODE
            _TESSERACT_RUNTIME_READY = False
            _TESSERACT_RUNTIME_DIR = None
            _TESSERACT_RUNTIME_MODE = ""

            tesseract_dir = candidate.parent.resolve()
            tessdata_dir = (tesseract_dir / "tessdata").resolve()

            if os.name == "nt":
                current_path = os.environ.get("PATH", "")
                prefix = str(tesseract_dir)
                if prefix and prefix.lower() not in current_path.lower():
                    os.environ["PATH"] = f"{prefix};{current_path}" if current_path else prefix
                add_dll_directory = getattr(os, "add_dll_directory", None)
                if callable(add_dll_directory):
                    try:
                        add_dll_directory(str(tesseract_dir))
                    except OSError:
                        pass

            if tessdata_dir.is_dir():
                # Το runtime probe θα αποφασίσει ποιος συνδυασμός env / --tessdata-dir δουλεύει πραγματικά.
                pass

            return str(candidate.resolve())

    raise RuntimeError(
        "Δεν βρέθηκε Tesseract OCR. "
        "Βάλε το bundled tesseract στον φάκελο third_party/tesseract "
        "ή κάνε system install."
    )


def ensure_tesseract_languages(required: list[str] | None = None) -> None:
    """Ελέγχει ότι ο Tesseract φορτώνει πραγματικά τις απαιτούμενες γλώσσες."""
    global _TESSERACT_RUNTIME_READY, _TESSERACT_RUNTIME_DIR, _TESSERACT_RUNTIME_MODE

    require_ocr_capability()

    required = required or ["ell", "eng"]

    if _TESSERACT_RUNTIME_READY:
        return

    tesseract_cmd = getattr(pytesseract.pytesseract, "tesseract_cmd", "") or configure_tesseract()
    runtime_dir, runtime_mode = _probe_tesseract_runtime(tesseract_cmd, required)

    _TESSERACT_RUNTIME_DIR = runtime_dir
    _TESSERACT_RUNTIME_MODE = runtime_mode
    _TESSERACT_RUNTIME_READY = True

    cmd_path = Path(str(tesseract_cmd)).resolve()
    tesseract_dir = cmd_path.parent

    if runtime_mode.startswith("env_tessdata"):
        _set_tessdata_prefix(tesseract_dir / "tessdata")
    elif runtime_mode.startswith("env_parent"):
        _set_tessdata_prefix(tesseract_dir)
    else:
        _set_tessdata_prefix(None)


def _run_tesseract_tsv(image: Image.Image, *, lang: str, base_config: str) -> pd.DataFrame:
    """
    Τρέχει απευθείας το native tesseract.exe σε προσωρινή PNG και επιστρέφει TSV σαν DataFrame.
    Αυτό αποφεύγει ιδιοτροπίες του pytesseract σε frozen Windows build.
    """
    require_ocr_capability()

    tesseract_cmd = configure_tesseract()
    ensure_tesseract_languages(lang.split("+"))

    cmd_path = Path(tesseract_cmd).resolve()
    cmd = [str(cmd_path)]

    with tempfile.TemporaryDirectory(prefix="ordermatcher_ocr_") as tmp_dir:
        tmp_path = Path(tmp_dir)
        input_path = tmp_path / "page.png"
        output_base = tmp_path / "ocr_output"
        tsv_path = tmp_path / "ocr_output.tsv"
        txt_path = tmp_path / "ocr_output.txt"

        image.save(input_path)

        cmd.extend([str(input_path), str(output_base), "-l", lang])

        if _TESSERACT_RUNTIME_DIR is not None:
            cmd.extend(["--tessdata-dir", _native_path(_TESSERACT_RUNTIME_DIR)])

        extra_args = [arg for arg in str(base_config).split() if arg]
        cmd.extend(extra_args)
        # Δεν βασιζόμαστε στο external config file "tsv" του Tesseract,
        # γιατί στο bundled build συνήθως πακετάρουμε μόνο traineddata και όχι
        # ολόκληρο το tessdata/configs. Με -c παράγεται TSV χωρίς εξάρτηση από configs.
        cmd.extend(["-c", "tessedit_create_tsv=1"])

        env = os.environ.copy()
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            env=env,
            errors="replace",
        )

        details = "\n".join(part for part in [result.stdout.strip(), result.stderr.strip()] if part).strip()

        if result.returncode != 0:
            raise RuntimeError(
                "Απέτυχε η native κλήση του bundled Tesseract. "
                f"Command: {' '.join(cmd)} | Έξοδος: {details}"
            )

        if not tsv_path.is_file():
            if txt_path.is_file():
                raise RuntimeError(
                    "Το Tesseract έτρεξε αλλά παρήγαγε TXT αντί για TSV. "
                    "Αυτό συνήθως σημαίνει ότι η TSV έξοδος ζητήθηκε με λάθος τρόπο ή μέσω missing config. "
                    f"Command: {' '.join(cmd)} | Έξοδος: {details}"
                )
            raise RuntimeError(
                "Το Tesseract ολοκλήρωσε χωρίς να παραχθεί TSV. "
                f"Command: {' '.join(cmd)} | Έξοδος: {details}"
            )

        return pd.read_csv(tsv_path, sep="\t")


def _safe_numeric_conf(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce").fillna(-1)


def _preprocess_ocr_variants(image: Image.Image) -> list[tuple[str, Image.Image]]:
    base = image.convert("L")
    variants: list[tuple[str, Image.Image]] = []
    clean = ImageOps.autocontrast(base)
    variants.append(("gray_autocontrast", clean))
    sharpened = clean.filter(ImageFilter.SHARPEN)
    binary = sharpened.point(lambda px: 255 if px > 170 else 0, mode="1").convert("L")
    variants.append(("binary_170", binary))
    denoised = clean.filter(ImageFilter.MedianFilter(size=3))
    binary_soft = ImageOps.autocontrast(denoised).point(lambda px: 255 if px > 150 else 0, mode="1").convert("L")
    variants.append(("denoise_binary_150", binary_soft))
    return variants


def _ocr_result_score(df: pd.DataFrame) -> float:
    if df is None or df.empty or "text" not in df.columns:
        return float("-inf")
    work = df.copy()
    work["text"] = work["text"].astype(str).str.strip()
    work = work[work["text"] != ""].copy()
    if work.empty:
        return float("-inf")
    conf = _safe_numeric_conf(work["conf"]) if "conf" in work.columns else pd.Series([-1] * len(work))
    tokens = work["text"].tolist()
    six_digit = sum(1 for t in tokens if len(safe_registry(t)) == 6)
    greekish = sum(1 for t in tokens if len(normalize_text(t)) >= 2 and re.search(r"[Α-ΩA-Z]", normalize_text(t)))
    avg_conf = float(conf[conf >= 0].mean()) if (conf >= 0).any() else 0.0
    low_conf_noise = int(((conf >= 0) & (conf < 20)).sum())
    return six_digit * 8 + greekish * 0.6 + avg_conf * 0.15 - low_conf_noise * 0.4


def _filter_ocr_tokens(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    work = df.copy()
    work["text"] = work["text"].astype(str).str.strip()
    work = work[work["text"] != ""].copy()
    if work.empty:
        return work
    conf = _safe_numeric_conf(work["conf"]) if "conf" in work.columns else pd.Series([-1] * len(work), index=work.index)
    norm = work["text"].map(normalize_text)
    registry_like = work["text"].map(safe_registry).str.len().ge(4)
    alpha_like = norm.str.contains(r"[Α-ΩA-Z]", regex=True)
    keep = registry_like | alpha_like | conf.ge(28) | conf.eq(-1)
    keep &= ~norm.str.fullmatch(r"[0-9]{1,2}")
    keep |= registry_like
    return work[keep].copy()


def normalize_text(value: object) -> str:
    """Κανονικοποιεί κείμενο ώστε να συγκρίνεται ανθεκτικά (τόνοι, πεζά/κεφαλαία, symbols)."""
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""

    text = str(value).strip().upper()
    text = " ".join(text.split())
    text = text.replace("Ϊ", "Ι").replace("Ϋ", "Υ")
    text = "".join(ch for ch in unicodedata.normalize("NFD", text) if unicodedata.category(ch) != "Mn")
    text = re.sub(r"[^A-ZΑ-Ω0-9 ]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def safe_registry(value: object) -> str:
    """Καθαρίζει το μητρώο από μορφές τύπου 12345.0, ΑΜ:123456 ή περίεργους separators."""
    if value is None:
        return ""
    if isinstance(value, float):
        if pd.isna(value):
            return ""
        if float(value).is_integer():
            return str(int(value))
    if isinstance(value, int):
        return str(value)

    text = str(value).strip()
    if not text:
        return ""

    if re.fullmatch(r"\d+(?:\.0+)?", text):
        return text.split(".", 1)[0]

    compact_digits = re.sub(r"\D", "", text)
    if len(compact_digits) >= 6:
        return compact_digits[:6]

    match = re.search(r"(\d{4,})", text)
    if match:
        return match.group(1)
    return compact_digits


def normalize_registry(value: object) -> str:
    reg = safe_registry(value)
    if not reg:
        return ""
    stripped = reg.lstrip("0")
    return stripped or "0"


def clean_patronymic(value: object) -> str:
    """Αφαιρεί το πρόθεμα ΤΟΥ / TOU και κρατά καθαρό το πατρώνυμο."""
    if value is None:
        return ""
    text = " ".join(str(value).split()).strip()
    text = re.sub(r"^(?:ΤΟΥ|TOU)\b[\s.\-–—]*", "", text, flags=re.IGNORECASE)
    return text.strip(" .-–—")



class PromotionPdfParser:
    """Διαβάζει διαταγές PDF είτε από native text layer είτε με OCR fallback."""
    SURNAME_X_MAX = 245
    NAME_X_MAX = 338
    PATRONYMIC_X_MAX = 468
    BASE_PAGE_WIDTH = 595
    OCR_ZOOM = 2.0
    OCR_CONFIGS = ("--oem 1 --psm 6", "--oem 1 --psm 4")

    NOISE_TOKENS = {
        "ΥΠ.",
        "ΓΡΑΦ.",
        "ΥΠ",
        "ΓΡΑΦ",
        "ΥΠ.ΓΡΑΦ.",
        "ΥΠ. ΓΡΑΦ.",
    }

    def __init__(self) -> None:
        self.last_mode = ""
        self.tesseract_path = ""
        self.last_ocr_variant = ""

    def parse(self, pdf_path: str | Path) -> pd.DataFrame:
        pdf_path = Path(pdf_path)

        try:
            df = self._parse_native(pdf_path)
            if not df.empty:
                self.last_mode = "Κανονική εξαγωγή κειμένου PDF"
                return df
        except Exception:
            pass

        if not is_ocr_build():
            raise RuntimeError(
                "Το PDF φαίνεται σκαναρισμένο ή χωρίς text layer. "
                "Χρησιμοποίησε την OCR έκδοση του προγράμματος."
            )

        df = self._parse_ocr(pdf_path)
        self.last_mode = "OCR σε σκαναρισμένο PDF"
        return df

    def _parse_native(self, pdf_path: Path) -> pd.DataFrame:
        rows: list[dict] = []

        with pdfplumber.open(str(pdf_path)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                words = page.extract_words(
                    x_tolerance=1,
                    y_tolerance=3,
                    keep_blank_chars=False,
                    use_text_flow=False,
                )
                grouped_by_line: dict[float, list[dict]] = defaultdict(list)
                for word in words:
                    grouped_by_line[round(word["top"], 1)].append(word)

                for top in sorted(grouped_by_line):
                    line_words = sorted(grouped_by_line[top], key=lambda item: item["x0"])
                    record = self._build_record_from_words(
                        line_words,
                        page_index=page_index,
                        surname_x_max=self.SURNAME_X_MAX,
                        name_x_max=self.NAME_X_MAX,
                        patronymic_x_max=self.PATRONYMIC_X_MAX,
                        x_key="x0",
                    )
                    if record:
                        rows.append(record)

        return self._finalize_dataframe(rows)

    def _parse_ocr(self, pdf_path: Path) -> pd.DataFrame:
        require_ocr_capability()

        if fitz is None:
            raise RuntimeError(
                "Η OCR έκδοση δεν έχει το package PyMuPDF. "
                "Τρέξε install στο OCR venv και ξανακάνε build."
            )

        self.tesseract_path = configure_tesseract()
        ensure_tesseract_languages(["ell", "eng"])

        rows: list[dict] = []
        selected_variants: list[str] = []

        with fitz.open(str(pdf_path)) as pdf:
            for page_index, page in enumerate(pdf, start=1):
                pixmap = page.get_pixmap(matrix=fitz.Matrix(self.OCR_ZOOM, self.OCR_ZOOM), alpha=False)
                image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
                page_width = image.width

                best_df = pd.DataFrame()
                best_score = float("-inf")
                best_variant_name = ""

                for variant_name, processed_image in _preprocess_ocr_variants(image):
                    for config in self.OCR_CONFIGS:
                        ocr_df = _run_tesseract_tsv(processed_image, lang="ell+eng", base_config=config)
                        ocr_df = _filter_ocr_tokens(ocr_df)
                        score = _ocr_result_score(ocr_df)
                        if score > best_score:
                            best_score = score
                            best_df = ocr_df
                            best_variant_name = f"{variant_name} | {config}"

                if best_df.empty:
                    continue

                selected_variants.append(f"σελ.{page_index}: {best_variant_name}")

                grouped_lines = [
                    group.sort_values("left")
                    for _, group in best_df.groupby(["block_num", "par_num", "line_num"], sort=True)
                ]

                surname_x_max, name_x_max, patronymic_x_max = self._estimate_column_boundaries(
                    grouped_lines,
                    page_width=page_width,
                )

                pending_record: Optional[dict] = None
                pending_top: Optional[int] = None

                for line_words_df in grouped_lines:
                    line_words = line_words_df.to_dict("records")
                    record = self._build_record_from_words(
                        line_words,
                        page_index=page_index,
                        surname_x_max=surname_x_max,
                        name_x_max=name_x_max,
                        patronymic_x_max=patronymic_x_max,
                        x_key="left",
                    )
                    current_top = int(line_words_df["top"].min())

                    if record:
                        rows.append(record)
                        pending_record = rows[-1]
                        pending_top = current_top
                        continue

                    if pending_record and pending_top is not None:
                        gap = current_top - pending_top
                        if self._looks_like_continuation(
                            line_words,
                            gap,
                            pending_record=pending_record,
                            name_x_max=name_x_max,
                            patronymic_x_max=patronymic_x_max,
                            x_key="left",
                        ):
                            self._append_continuation(
                                pending_record,
                                line_words,
                                surname_x_max=surname_x_max,
                                name_x_max=name_x_max,
                                patronymic_x_max=patronymic_x_max,
                                x_key="left",
                            )
                            pending_top = current_top
                        else:
                            pending_record = None
                            pending_top = None

        self.last_ocr_variant = " ; ".join(selected_variants[:6])
        return self._finalize_dataframe(rows)

    def _estimate_column_boundaries(
        self,
        grouped_lines: list[pd.DataFrame],
        *,
        page_width: int,
    ) -> tuple[float, float, float]:
        surname_fallback = page_width * (self.SURNAME_X_MAX / self.BASE_PAGE_WIDTH)
        name_fallback = page_width * (self.NAME_X_MAX / self.BASE_PAGE_WIDTH)
        patronymic_fallback = page_width * (self.PATRONYMIC_X_MAX / self.BASE_PAGE_WIDTH)

        surname_starts: list[float] = []
        name_starts: list[float] = []
        patronymic_starts: list[float] = []

        for line_df in grouped_lines:
            words = line_df.to_dict("records")
            parsed = self._extract_registry_and_start_index(words)
            if parsed is None:
                continue
            _, start_index = parsed
            content = words[start_index:]
            if len(content) < 2:
                continue

            columns: list[float] = []
            current_bucket = [float(content[0]["left"])]
            prev_left = float(content[0]["left"])
            for word in content[1:]:
                current_left = float(word["left"])
                if current_left - prev_left > page_width * 0.045:
                    columns.append(min(current_bucket))
                    current_bucket = [current_left]
                else:
                    current_bucket.append(current_left)
                prev_left = current_left
            columns.append(min(current_bucket))

            if len(columns) >= 1:
                surname_starts.append(columns[0])
            if len(columns) >= 2:
                name_starts.append(columns[1])
            if len(columns) >= 3:
                patronymic_starts.append(columns[2])

        surname_cut = (
            (float(pd.Series(surname_starts).median()) + float(pd.Series(name_starts).median())) / 2
            if surname_starts and name_starts
            else surname_fallback
        )
        name_cut = (
            (float(pd.Series(name_starts).median()) + float(pd.Series(patronymic_starts).median())) / 2
            if name_starts and patronymic_starts
            else name_fallback
        )
        patronymic_cut = (
            min(patronymic_fallback, float(pd.Series(patronymic_starts).median()) + page_width * 0.09)
            if patronymic_starts
            else patronymic_fallback
        )

        surname_cut = max(page_width * 0.22, min(surname_cut, page_width * 0.50))
        name_cut = max(surname_cut + page_width * 0.08, min(name_cut, page_width * 0.70))
        patronymic_cut = max(name_cut + page_width * 0.08, min(patronymic_cut, page_width * 0.88))
        return surname_cut, name_cut, patronymic_cut

    def _extract_registry_and_start_index(self, line_words: list[dict]) -> Optional[tuple[str, int]]:
        texts = [str(word.get("text", "")).strip() for word in line_words if str(word.get("text", "")).strip()]
        if len(texts) < 2:
            return None

        for idx in range(min(3, len(texts))):
            direct = safe_registry(texts[idx])
            if len(direct) == 6:
                return direct, idx + 1

            if idx + 1 < len(texts):
                merged = safe_registry(f"{texts[idx]}{texts[idx + 1]}")
                if len(merged) == 6:
                    return merged, idx + 2

        return None

    def _build_record_from_words(
        self,
        line_words: list[dict],
        *,
        page_index: int,
        surname_x_max: float,
        name_x_max: float,
        patronymic_x_max: float,
        x_key: str,
    ) -> Optional[dict]:
        texts = [str(word.get("text", "")).strip() for word in line_words if str(word.get("text", "")).strip()]
        if len(texts) < 2:
            return None

        row_no: Optional[int] = None
        row_digits = re.sub(r"\D", "", texts[0])
        if row_digits and len(row_digits) <= 3:
            row_no = int(row_digits)

        parsed = self._extract_registry_and_start_index(line_words)
        if parsed is None:
            return None

        registry, start_index = parsed
        if start_index >= len(line_words):
            return None

        record = {
            "pdf_page": page_index,
            "pdf_row_no": row_no,
            "registry": registry,
            "surname": [],
            "name": [],
            "patronymic": [],
            "extra": [],
        }

        for word in line_words[start_index:]:
            x = float(word[x_key])
            token = str(word.get("text", "")).strip()
            if not token:
                continue

            conf_raw = str(word.get("conf", "")).strip()
            conf = float(conf_raw) if conf_raw not in {"", "nan", "None"} else -1
            normalized = normalize_text(token)
            if conf >= 0 and conf < 15 and not normalized and not safe_registry(token):
                continue

            if x < surname_x_max:
                record["surname"].append(token)
            elif x < name_x_max:
                record["name"].append(token)
            elif x < patronymic_x_max:
                record["patronymic"].append(token)
            else:
                record["extra"].append(token)

        return self._record_to_output(record)

    def _record_to_output(self, record: dict) -> Optional[dict]:
        surname = " ".join(record["surname"]).strip()
        name = " ".join(record["name"]).strip()
        patronymic_tokens = [t for t in record["patronymic"] if normalize_text(t) not in self.NOISE_TOKENS]
        extra_tokens = [t for t in record["extra"] if normalize_text(t) not in self.NOISE_TOKENS]
        patronymic = clean_patronymic(" ".join(patronymic_tokens))
        extra = " ".join(extra_tokens).strip()

        if not surname or not name:
            return None

        return {
            "pdf_page": record["pdf_page"],
            "pdf_row_no": record["pdf_row_no"],
            "registry": record["registry"],
            "surname": surname,
            "name": name,
            "patronymic": patronymic,
            "extra": extra,
        }

    def _looks_like_continuation(
        self,
        line_words: list[dict],
        gap: int,
        *,
        pending_record: Optional[dict] = None,
        name_x_max: Optional[float] = None,
        patronymic_x_max: Optional[float] = None,
        x_key: str = "left",
    ) -> bool:
        if not line_words or gap > 65:
            return False

        texts = [str(word.get("text", "")).strip() for word in line_words if str(word.get("text", "")).strip()]
        if not texts or len(texts) > 5:
            return False

        if self._extract_registry_and_start_index(line_words) is not None:
            return False

        normalized = [normalize_text(t) for t in texts]
        if not any(re.search(r"[Α-ΩA-Z]", token) for token in normalized):
            return False

        if pending_record and not pending_record.get("patronymic") and patronymic_x_max is not None:
            lefts = [float(word.get(x_key, 0)) for word in line_words]
            if lefts and min(lefts) >= (name_x_max or 0) * 0.95:
                return True

        return True

    def _append_continuation(
        self,
        record: dict,
        line_words: list[dict],
        *,
        surname_x_max: float,
        name_x_max: float,
        patronymic_x_max: float,
        x_key: str,
    ) -> None:
        for word in line_words:
            x = float(word[x_key])
            token = str(word.get("text", "")).strip()
            if not token:
                continue

            if x < surname_x_max:
                record["surname"] = (f"{record.get('surname', '')} {token}").strip()
            elif x < name_x_max:
                record["name"] = (f"{record.get('name', '')} {token}").strip()
            elif x < patronymic_x_max:
                current = record.get("patronymic", "")
                record["patronymic"] = clean_patronymic(f"{current} {token}")
            else:
                record["extra"] = (f"{record.get('extra', '')} {token}").strip()

    def _finalize_dataframe(self, rows: list[dict]) -> pd.DataFrame:
        df = pd.DataFrame(rows)
        if df.empty:
            raise ValueError("Δεν βρέθηκαν εγγραφές στο PDF.")

        df = df.drop_duplicates(subset=["registry", "surname", "name", "patronymic"]).reset_index(drop=True)
        df["norm_registry"] = df["registry"].map(normalize_registry)
        for col in ["surname", "name", "patronymic"]:
            df[f"norm_{col}"] = df[col].map(normalize_text)
        return df


class ServiceExcelLoader:
    """Φορτώνει το Excel δύναμης και εντοπίζει αυτόματα τις σωστές στήλες."""
    REQUIRED_COLUMNS = {"ΜΗΤΡΩΟ", "ΕΠΩΝΥΜΟ", "ΟΝΟΜΑ"}

    def load(self, excel_path: str | Path) -> pd.DataFrame:
        """Δοκιμάζει τον κατάλληλο engine ανάλογα με το αν το αρχείο είναι xls ή xlsx."""
        best_error = None
        excel_path = Path(excel_path)

        engines: list[Optional[str]] = [None]
        if excel_path.suffix.lower() == ".xls":
            engines.append("xlrd")
        elif excel_path.suffix.lower() == ".xlsx":
            engines.append("openpyxl")

        xl = None
        chosen_engine: Optional[str] = None
        for engine in engines:
            try:
                xl = pd.ExcelFile(excel_path, engine=engine)
                chosen_engine = engine
                break
            except Exception as exc:  # noqa: BLE001
                best_error = exc

        if xl is None:
            if best_error:
                raise best_error
            raise ValueError("Δεν ήταν δυνατή η φόρτωση του Excel.")

        for sheet_name in xl.sheet_names:
            try:
                return self._load_sheet(excel_path, sheet_name, engine=chosen_engine)
            except Exception as exc:  # noqa: BLE001
                best_error = exc
        if best_error:
            raise best_error
        raise ValueError("Δεν ήταν δυνατή η φόρτωση του Excel.")

    def _load_sheet(self, excel_path: str | Path, sheet_name: str, engine: Optional[str] = None) -> pd.DataFrame:
        raw = pd.read_excel(excel_path, sheet_name=sheet_name, header=None, engine=engine)
        header_row_index = self._find_header_row(raw)
        if header_row_index is None:
            raise ValueError(f"Δεν βρέθηκε γραμμή επικεφαλίδων στο φύλλο: {sheet_name}")

        header_values = [normalize_text(v) for v in raw.iloc[header_row_index].tolist()]
        df = raw.iloc[header_row_index + 1 :].copy()
        df.columns = header_values
        df = df.dropna(how="all").reset_index(drop=True)

        keep_cols = [c for c in df.columns if c]
        df = df[keep_cols]

        registry_col = self._pick_column(df.columns, ["ΜΗΤΡΩΟ", "ΑΡΙΘΜΟΣ ΜΗΤΡΩΟΥ", "Α Μ", "ΑΜ"])
        surname_col = self._pick_column(df.columns, ["ΕΠΩΝΥΜΟ"])
        name_col = self._pick_column(df.columns, ["ΟΝΟΜΑ"])
        patronymic_col = self._pick_column(df.columns, ["ΠΑΤΡΩΝΥΜΟ"])
        rank_col = self._pick_column(df.columns, ["ΒΑΘΜΟΣ"]) or ""
        unit_col = self._pick_column(df.columns, ["ΟΡΓΑΝΙΚΗ", "ΥΠΗΡΕΣΙΑ", "ΤΜΗΜΑ", "ΔΙΕΥΘΥΝΣΗ"]) or ""

        if not registry_col or not surname_col or not name_col:
            raise ValueError("Δεν εντοπίστηκαν οι βασικές στήλες Μητρώο / Επώνυμο / Όνομα στο Excel.")

        out = pd.DataFrame(
            {
                "registry": df[registry_col].map(safe_registry),
                "rank": df[rank_col].astype(str).replace("nan", "").str.strip() if rank_col else "",
                "surname": df[surname_col].astype(str).replace("nan", "").str.strip(),
                "name": df[name_col].astype(str).replace("nan", "").str.strip(),
                "patronymic": df[patronymic_col].astype(str).replace("nan", "").str.strip() if patronymic_col else "",
                "service_unit": df[unit_col].astype(str).replace("nan", "").str.strip() if unit_col else "",
                "source_sheet": sheet_name,
            }
        )
        out = out[out["registry"].astype(str).str.strip() != ""].copy()
        out = out.reset_index(drop=True)

        out["norm_registry"] = out["registry"].map(normalize_registry)
        for col in ["surname", "name", "patronymic"]:
            out[f"norm_{col}"] = out[col].map(normalize_text)
        return out

    def _find_header_row(self, raw: pd.DataFrame) -> Optional[int]:
        for idx, row in raw.iterrows():
            normalized = {normalize_text(v) for v in row.tolist()}
            if self.REQUIRED_COLUMNS.issubset(normalized):
                return int(idx)
        return None

    @staticmethod
    def _pick_column(columns: list[str] | pd.Index, candidates: list[str]) -> Optional[str]:
        normalized_columns = [normalize_text(c) for c in columns]
        for candidate in candidates:
            target = normalize_text(candidate)
            for col, norm_col in zip(columns, normalized_columns):
                if norm_col == target:
                    return str(col)
        for candidate in candidates:
            target = normalize_text(candidate)
            for col, norm_col in zip(columns, normalized_columns):
                if target and target in norm_col:
                    return str(col)
        return None


class MatcherEngine:
    """Εκτελεί τη σύγκριση PDF ↔ Excel με τρία επίπεδα ταύτισης."""
    def match(self, promotions_df: pd.DataFrame, service_df: pd.DataFrame) -> tuple[pd.DataFrame, dict]:
        # 1) Ταύτιση με μητρώο
        # 2) Ταύτιση με ονοματεπώνυμο + πατρώνυμο
        # 3) Ταύτιση μόνο με ονοματεπώνυμο όταν είναι μοναδικό και στις δύο πλευρές
        if promotions_df.empty:
            raise ValueError("Το PDF δεν έδωσε εγγραφές για σύγκριση.")
        if service_df.empty:
            raise ValueError("Το Excel δεν έδωσε εγγραφές για σύγκριση.")

        promotions_df = promotions_df.copy().reset_index(drop=True)
        service_df = service_df.copy().reset_index(drop=True)
        promotions_df["_pdf_id"] = promotions_df.index.astype(int)
        service_df["_excel_id"] = service_df.index.astype(int)

        exact_registry = promotions_df.merge(
            service_df,
            left_on="norm_registry",
            right_on="norm_registry",
            how="inner",
            suffixes=("_pdf", "_excel"),
        ).copy()
        if not exact_registry.empty:
            exact_registry["match_method"] = "Μητρώο"
            exact_registry["name_key"] = (
                exact_registry["norm_surname_pdf"]
                + "|"
                + exact_registry["norm_name_pdf"]
                + "|"
                + exact_registry["norm_patronymic_pdf"]
            )

        used_pdf_ids = set(exact_registry.get("_pdf_id", pd.Series(dtype=int)).tolist())
        used_excel_ids = set(exact_registry.get("_excel_id", pd.Series(dtype=int)).tolist())

        pdf_left = promotions_df[~promotions_df["_pdf_id"].isin(used_pdf_ids)].copy()
        excel_left = service_df[~service_df["_excel_id"].isin(used_excel_ids)].copy()

        pdf_left["name_key"] = (
            pdf_left["norm_surname"] + "|" + pdf_left["norm_name"] + "|" + pdf_left["norm_patronymic"]
        )
        excel_left["name_key"] = (
            excel_left["norm_surname"] + "|" + excel_left["norm_name"] + "|" + excel_left["norm_patronymic"]
        )

        exact_names = pdf_left.merge(
            excel_left,
            on="name_key",
            how="inner",
            suffixes=("_pdf", "_excel"),
        ).copy()
        if not exact_names.empty:
            exact_names["match_method"] = "Ονοματεπώνυμο + Πατρώνυμο"
            exact_names["norm_registry"] = exact_names["norm_registry_pdf"]

        used_pdf_ids.update(exact_names.get("_pdf_id", pd.Series(dtype=int)).tolist())
        used_excel_ids.update(exact_names.get("_excel_id", pd.Series(dtype=int)).tolist())

        pdf_left = promotions_df[~promotions_df["_pdf_id"].isin(used_pdf_ids)].copy()
        excel_left = service_df[~service_df["_excel_id"].isin(used_excel_ids)].copy()

        pdf_left["name_key_basic"] = pdf_left["norm_surname"] + "|" + pdf_left["norm_name"]
        excel_left["name_key_basic"] = excel_left["norm_surname"] + "|" + excel_left["norm_name"]

        pdf_unique_basic = pdf_left["name_key_basic"].value_counts()
        excel_unique_basic = excel_left["name_key_basic"].value_counts()
        pdf_left = pdf_left[pdf_left["name_key_basic"].map(pdf_unique_basic).eq(1)].copy()
        excel_left = excel_left[excel_left["name_key_basic"].map(excel_unique_basic).eq(1)].copy()

        exact_basic = pdf_left.merge(
            excel_left,
            on="name_key_basic",
            how="inner",
            suffixes=("_pdf", "_excel"),
        ).copy()
        if not exact_basic.empty:
            exact_basic["name_key"] = exact_basic["name_key_basic"]
            exact_basic["match_method"] = "Ονοματεπώνυμο"
            exact_basic["norm_registry"] = exact_basic["norm_registry_pdf"]

        common = pd.concat([exact_registry, exact_names, exact_basic], ignore_index=True, sort=False)
        if not common.empty:
            common = common.drop_duplicates(subset=["_pdf_id", "_excel_id"], keep="first")

        result = pd.DataFrame(
            {
                "Α/Α PDF": common.get("pdf_row_no", ""),
                "Μητρώο": common.get("registry_pdf", common.get("registry_excel", "")),
                "Επώνυμο": common.get("surname_pdf", common.get("surname_excel", "")),
                "Όνομα": common.get("name_pdf", common.get("name_excel", "")),
                "Πατρώνυμο": common.get("patronymic_pdf", common.get("patronymic_excel", "")),
                "Βαθμός": common.get("rank", ""),
                "Οργανική": common.get("service_unit", ""),
                "Τρόπος Ταύτισης": common.get("match_method", ""),
                "PDF Σελίδα": common.get("pdf_page", ""),
                "Excel Φύλλο": common.get("source_sheet", ""),
            }
        )
        result = result.sort_values(by=["Α/Α PDF", "Επώνυμο", "Όνομα"], na_position="last").reset_index(drop=True)

        matched_pdf_ids = set(common.get("_pdf_id", pd.Series(dtype=int)).tolist())
        matched_excel_ids = set(common.get("_excel_id", pd.Series(dtype=int)).tolist())
        only_promotions = promotions_df[~promotions_df["_pdf_id"].isin(matched_pdf_ids)].copy()
        only_service = service_df[~service_df["_excel_id"].isin(matched_excel_ids)].copy()

        summary = {
            "promotions_total": int(len(promotions_df)),
            "service_total": int(len(service_df)),
            "common_total": int(len(result)),
            "registry_matches": int((result["Τρόπος Ταύτισης"] == "Μητρώο").sum()) if not result.empty else 0,
            "name_patronymic_matches": int((result["Τρόπος Ταύτισης"] == "Ονοματεπώνυμο + Πατρώνυμο").sum()) if not result.empty else 0,
            "name_only_matches": int((result["Τρόπος Ταύτισης"] == "Ονοματεπώνυμο").sum()) if not result.empty else 0,
            "only_promotions_total": int(len(only_promotions)),
            "only_service_total": int(len(only_service)),
        }

        return result, {
            "common": result,
            "only_promotions": only_promotions,
            "only_service": only_service,
            "summary": summary,
        }


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


class WordExporter:
    """Εξάγει τα κοινά πρόσωπα και τη σύνοψη σε ευανάγνωστο Word report."""
    def export(self, output_path: str | Path, common_df: pd.DataFrame, summary: dict, source_pdf: str, source_excel: str) -> None:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.3)
        section.left_margin = Cm(1.05)
        section.right_margin = Cm(1.05)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        styles = doc.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(8.5)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(3)
        run = title.add_run("Κοινοί μεταξύ Διαταγής και Δύναμης Υπηρεσίας")
        run.bold = True
        run.font.size = Pt(15)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(9)
        subtitle_run = subtitle.add_run(datetime.now().strftime("Ημερομηνία εξαγωγής: %d/%m/%Y %H:%M"))
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(8.5)

        meta = doc.add_table(rows=4, cols=2)
        meta.style = "Table Grid"
        meta.autofit = False
        set_table_fixed_layout(meta)
        meta_widths = [Cm(4.1), Cm(22.0)]
        meta_rows = [
            ("PDF Διαταγής", str(source_pdf)),
            ("Excel Δύναμης", str(source_excel)),
            ("Σύνολο κοινών", str(summary.get("common_total", 0))),
            (
                "Τρόπος ταύτισης",
                f"Μητρώο: {summary.get('registry_matches', 0)} | Ονοματεπώνυμο + Πατρώνυμο: {summary.get('name_patronymic_matches', 0)} | Ονοματεπώνυμο: {summary.get('name_only_matches', 0)}",
            ),
        ]
        for i, (left, right) in enumerate(meta_rows):
            for j, value in enumerate((left, right)):
                cell = meta.cell(i, j)
                cell.width = meta_widths[j]
                cell.text = str(value)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell, top=55, start=95, bottom=55, end=95)
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                if paragraph.runs:
                    paragraph.runs[0].font.size = Pt(8.5)
            if meta.cell(i, 0).paragraphs[0].runs:
                meta.cell(i, 0).paragraphs[0].runs[0].bold = True

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(5)

        table = doc.add_table(rows=1, cols=len(TABLE_COLUMNS))
        table.style = "Table Grid"
        table.autofit = False
        set_table_fixed_layout(table)
        widths_cm = [1.05, 1.55, 3.35, 3.05, 2.8, 1.55, 5.6, 4.1]
        header_row = table.rows[0]
        for idx, col_name in enumerate(TABLE_COLUMNS):
            cell = header_row.cells[idx]
            cell.width = Cm(widths_cm[idx])
            cell.text = col_name
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=65, start=85, bottom=65, end=85)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            if paragraph.runs:
                paragraph.runs[0].bold = True
                paragraph.runs[0].font.size = Pt(8)
        set_repeat_table_header(header_row)

        if common_df.empty:
            row = table.add_row()
            row.cells[0].merge(row.cells[-1])
            row.cells[0].text = "Δεν βρέθηκαν κοινά άτομα."
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_margins(row.cells[0], top=70, start=100, bottom=70, end=100)
        else:
            for _, item in common_df.iterrows():
                row = table.add_row()
                values = [
                    item.get("Α/Α PDF", ""),
                    item.get("Μητρώο", ""),
                    item.get("Επώνυμο", ""),
                    item.get("Όνομα", ""),
                    item.get("Πατρώνυμο", ""),
                    item.get("Βαθμός", ""),
                    item.get("Οργανική", ""),
                    item.get("Τρόπος Ταύτισης", ""),
                ]
                for idx, value in enumerate(values):
                    cell = row.cells[idx]
                    cell.width = Cm(widths_cm[idx])
                    cell.text = "" if pd.isna(value) else str(value)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    set_cell_margins(cell, top=55, start=85, bottom=55, end=85)
                    paragraph = cell.paragraphs[0]
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1, 5) else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        doc.save(str(output_path))
